
import docx

file_inv = open('guests.txt', 'r')
invitations = file_inv.read().splitlines()

# Open this document to use pre-set styles (my_heading, my_normal, my_bold)
template_doc = docx.Document('invitations_template.docx')


for invitation in invitations:
    para_obj0 = template_doc.add_paragraph('It would be a pleasure to have the company of', 'my_heading')
    para_obj1 = template_doc.add_paragraph(invitation, 'my_bold')
    para_obj2 = template_doc.add_paragraph('at 11010 Memory Lane on the evening of', 'my_normal')
    para_obj3 = template_doc.add_paragraph('April 1st', 'my_normal')
    para_obj4 = template_doc.add_paragraph('at 7 o\'clock', 'my_normal')
    template_doc.add_picture('zophie.png', width=docx.shared.Cm(6), style='my_normal')
    template_doc.add_page_break()

template_doc.save(f'invitations.docx')
